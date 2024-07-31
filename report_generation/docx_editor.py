from docx import Document
import os
from docx.oxml import parse_xml
import re
from report_generation.interaction import Interaction
from docx.shared import Inches

class DocxEditor:
    """
    A class to edit a .docx file.
    """

    def __init__(self, file_path : str):
        """
        Initialize the DocxEditor with a .docx file path.
        
        :param file_path: The path to the .docx file to be edited.
        """
        self.file_path = file_path
        self.doc = Document(file_path)

    def add_text(self, term : str, text_to_add : str):
        """
        Add a new paragraph with the specified text after the paragraph containing the term.
        
        :param term: The term to search for in the document.
        :param text_to_add: The text to add in a new paragraph after the paragraph containing the term.
        """
        for para in self.doc.paragraphs:
            if term in para.text:
                new_para_xml = parse_xml(
                    r'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:p>')
                para._p.addnext(new_para_xml)
                new_run = parse_xml(
                    r'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:t xml:space="preserve">%s</w:t></w:r>' % text_to_add)
                new_para_xml.append(new_run)
                break

    def add_bullet_list(self, term : str, bullet_list : list):
        """
        Add a bullet list after the paragraph containing the term.
        
        :param term: The term to search for in the document.
        :param bullet_list: The list of items to add as bullet points after the paragraph containing the term.
        """
        bullet_list.reverse()
        for para in self.doc.paragraphs:
            if term in para.text:
                for item in bullet_list:
                    new_para_xml = parse_xml(
                        r'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:p>')
                    para._p.addnext(new_para_xml)
                    new_run = parse_xml(
                        r'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:t xml:space="preserve">• %s</w:t></w:r>' % item)
                    new_para_xml.append(new_run)
                break

    def replace_text_block(self, term : str, replacement_text: str):
        """
        Replace the paragraph containing the term with the replacement text.
        
        :param term: The term to search for in the document.
        :param replacement_text: The text to replace the paragraph containing the term.
        """
        for para in self.doc.paragraphs:
            if term in para.text:
                for run in para.runs[::-1]:
                    run.clear()
                para.add_run(replacement_text)
                break

    def save_changes(self, new_file_name=None):
        """
        Save the changes to a new file.
        If a new file name is provided, use it. Otherwise, use the default name 'new_' + old file name.
        
        :param new_file_name: The name of the new file where the changes will be saved. If not provided, the default name 'new_' + old file name will be used.
        """
        dir_name = os.path.dirname(self.file_path)
        if new_file_name is None:
            base_name = 'new_' + os.path.basename(self.file_path)
        else:
            base_name = new_file_name + os.path.splitext(self.file_path)[1]
        self.new_file_path = os.path.join(dir_name, base_name)
        self.doc.save(self.new_file_path)

   
    def replace_specific_text(self, para, term: str, replacement_text: str):
        inline = para.runs
        started = False
        key_index = 0
        found_runs = []
        found_all = False
        replace_done = False

        for i in range(len(inline)):
            run_text = inline[i].text
            
            # Case 1: Term is entirely within one run
            if term in run_text and not started:
                new_text = run_text.replace(term, replacement_text)
                inline[i].text = new_text
                replace_done = True
                found_all = True            
        
        for i in range(len(inline)):
            run_text = inline[i].text
            found_all = False
            replace_done = False
            
            # Avoid empty string access
            if term[key_index] not in run_text and not started:
                # Continue looking for the term
                continue

            # Case 2: Search for partial text, find the first run
            if term[key_index] in run_text and (len(run_text) > 0 and run_text[-1] in term) and not started:
                start_index = run_text.find(term[key_index])
                check_length = len(run_text)
                for text_index in range(start_index, check_length):
                    if run_text[text_index] != term[key_index]:
                        # No match, so break
                        break
                if key_index == 0:
                    started = True
                chars_found = check_length - start_index
                key_index += chars_found
                found_runs.append((i, start_index, chars_found))
                if key_index < len(term):
                    continue
                else:
                    found_all = True
                    break

            # Case 3: Search for partial text, find subsequent runs
            if term[key_index] in run_text and started and not found_all:
                chars_found = 0
                check_length = len(run_text)
                for text_index in range(0, check_length):
                    if run_text[text_index] == term[key_index]:
                        key_index += 1
                        chars_found += 1
                    else:
                        break
                found_runs.append((i, 0, chars_found))
                if key_index == len(term):
                    found_all = True
                    break

        if found_all and not replace_done:
            for i, item in enumerate(found_runs):
                index, start, length = item
                run_text = inline[index].text
                if i == 0:
                    # Replace the term with replacement_text in the first run
                    text = run_text.replace(run_text[start:start + length], replacement_text)
                    inline[index].text = text
                else:
                    # Remove parts of the term from subsequent runs
                    text = run_text.replace(run_text[start:start + length], '')
                    inline[index].text = text
    

    def replace_paragraphs(self, para, data, sheet_name):
        pattern = r"Excel[a-zA-Z]{1,2}\d+"
        matches = re.findall(pattern, para.text)
        for match in matches:
            cell_reference = match[5:]
            replacement_text = str(data.get_data_at_cell(sheet_name, cell_reference))
            self.replace_specific_text(para, match, replacement_text)

    def replace_key_words_in_part(self, part, data, sheet_name):
        for para in part.paragraphs:
            self.replace_paragraphs(para, data, sheet_name)

    def replace_key_words_in_table(self, table, data, sheet_name):
        for row in table.rows:
            for cell in row.cells:
                self.replace_key_words_in_part(cell, data, sheet_name)

    def replace_key_words(self, data, sheet_name, key_word=None):
        """
        This function replaces key words in the document with data from an Excel sheet. It searches for patterns of the form "Excel (word)" in the document. For each match, it retrieves the corresponding data from the specified Excel sheet and replaces the match with this data.

        :param data: An object which contains the data from the Excel sheet.
        :param sheet_name: The name of the Excel sheet from which to retrieve data.
        :param key_word: The key word to search for in the document. If not provided, the function will search for the pattern "Excel (word)".
        """
        # Replace in the main body of the document
        for para in self.doc.paragraphs:
            self.replace_paragraphs(para, data, sheet_name)
        
        # Replace in headers and footers
        for section in self.doc.sections:
            self.replace_key_words_in_part(section.header, data, sheet_name)
            self.replace_key_words_in_part(section.footer, data, sheet_name) 
        # Replace in tables
        for table in self.doc.tables:
            self.replace_key_words_in_table(table, data, sheet_name)         

    def replace_term_with_image(self, term: str, folder_path: str):
        """
        Search for the term in the document and replace it with an image from the specified file path.
        The image is scaled to fit the width of the page and a caption is added below the image.
        The caption is formatted with the "Caption" style.

        :param term: The term to search for in the document.
        """
        for i, para in enumerate(self.doc.paragraphs):
            if term in para.text:
                file_name, img_cap = self.get_image_data(para.text)
                img_path = self.find_image(file_name, folder_path)
                if img_path != None:

                    para.clear()
                    run = para.add_run()
                    run.add_picture(img_path, width=Inches(6))
                    # Add the caption to a new paragraph with the "Caption" style
                    new_para_xml = parse_xml(r'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:pStyle w:val="Caption"/></w:pPr></w:p>')
                    para._p.addnext(new_para_xml)
                    new_run = parse_xml(r'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:t xml:space="preserve">%s</w:t></w:r>' % img_cap)
                    new_para_xml.append(new_run)
                    #self.doc.paragraphs[i+1].style = self.doc.styles['Caption']


    @staticmethod
    def get_image_data(info : str):
        """
        This function extracts the image name and caption from a string. The string is expected to be in the format "##Image##image_name##image_caption##".

        :param info: The string containing the image name and caption.
        :return: A tuple containing the image name and caption.
        """
        parts = info.split("##")

        img_name = parts[2]
        img_caption = parts[3]

        return img_name, img_caption

    @staticmethod
    def find_image(image_name :str, folder_path : str):
        """
        This function searches for an image file in a specified folder. The image file is expected to have a specific name and be in one of the
        following formats: .jpeg, .jpg, .png, .tiff.

        :param image_name: The name of the image file to search for.
        :param folder_path: The path of the folder in which to search for the image file.
        :return: The path of the image file if found, otherwise None.
        """
        image_formats = [".jpeg", ".jpg", ".png",".tiff"]

        for filename in os.listdir(folder_path):
            if filename.startswith(image_name) and any(filename.endswith(format) for format in image_formats):
                return os.path.join(folder_path, filename)

        return None
