import docx
from docx.shared import Inches

class ReportComponents:
    """
    A class to add elements to a Word report.

    """

    def __init__(self, project_data):
        """Constructs all the necessary attributes for the report object.

        Parameters
        ----------
            project_data : dict
                project data to be included in the report
        """
        self.doc = docx.Document()
        self.project_data = project_data

    def add_section(self, title):
        """Adds a section with the given title to the document."""
        self.doc.add_heading(title, level=1)

    def add_header_and_footer(self, header, footer):
        """Adds a header and footer to the document."""
        section = self.doc.sections[0]
        section.header.paragraphs[0].text = header
        section.footer.paragraphs[0].text = footer

    def add_text(self, text):
        """Adds a paragraph of text to the document."""
        self.doc.add_paragraph(text)

    def add_bullet_list(self, items):
        """Adds a bullet list to the document."""
        for item in items:
            self.doc.add_paragraph(item, style='ListBullet')

    def add_image(self, img_path, width=Inches(1.25)):
        """Adds an image to the document."""
        self.doc.add_picture(img_path, width=width)

    def add_table(self, data):
        """Adds a table to the document."""
        table = self.doc.add_table(rows=1, cols=len(data[0]))
        for item in data:
            cells = table.add_row().cells
            for i in range(len(item)):
                cells[i].text = str(item[i])

    def save_report(self, filename):
        """Saves the document to a file."""
        self.doc.save(filename)
